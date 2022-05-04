import pyscreenshot as ImageGrab
import os
from tkinter import *
from tkinter import filedialog
import pyautogui as pygi
import docx
from docx.shared import Inches
import aspose.words as aw





# window_initial
root  = Tk()
root.title('螢幕擷取小幫手')
root.geometry('250x200')

folder_path = StringVar()
num = IntVar()

# function
def screen_shot(): #傳入位置
  num.set(num.get()+ 1)   
  im = ImageGrab.grab(bbox=(193, 255, 1100,   500))
  saving_name = folder_path.get() + '/image' + str(num.get()) + '.png'
  im.save(saving_name)
  print(saving_name, path.exists(saving_name))

def select_path():
  global folder_path
  filename = filedialog.askdirectory()
  print("open folder: " + filename)
  folder_path.set(filename)
  
# 匯出成pdf 並刪除全部png
def output(): # 參數: 路徑& 照片數
  document = docx.Document()
  for i in range(1,num.get()+1):
    document.add_picture(folder_path.get() + '\\image' + str(i) + '.png',  width=Inches(6.0))
  document.save(folder_path.get() + '\\''sheet.docx')
  # docx to pdf
  doc = aw.Document(folder_path.get() + '\\'"sheet.docx")
  doc.save(folder_path.get() + '\\'"sheet.pdf")
  
  #迴圈delete all photos
  for i in range(1,num.get()+1):
    os.remove(folder_path.get() + '\\image' + str(i) + '.png')
  #歸零
  num.set(0)




# 選取路徑
button1 = Button(text="選擇資料夾", command = select_path)
button1.grid(row=0, sticky="w")


# 顯示當前路徑
label1 = Label(textvariable= folder_path)  
label1.grid(row=0, column=1)

# 螢幕截圖按鈕
button2 = Button(root, text='screen_shot', command = screen_shot,
                 height=5, width=10, bg='#c92036')
button2.grid(row=1,sticky="w")

# 截圖次數顯示
label2 = Label(textvariable= num)
label2.grid(row=2, sticky="w")


# 匯出檔案
button3 = Button(root, text="匯出", command= output)
button3.grid(row=3, sticky="w")


root.mainloop()


